"""
各种文件类型的解析器
"""
import re
import json
from typing import Dict, Any, List
from pathlib import Path
import aiofiles


async def parse_chat_records(file_path: str) -> Dict[str, Any]:
    """
    解析聊天记录文件
    支持微信、QQ 等格式
    """
    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = await f.read()
    
    segments = []
    
    # 尝试识别格式
    # 微信格式: [时间] 昵称: 消息
    wechat_pattern = r'\[(\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2})\]\s+([^:]+):\s+(.+)'
    wechat_matches = re.findall(wechat_pattern, content)
    
    if wechat_matches:
        # 微信格式
        for timestamp, sender, message in wechat_matches:
            segments.append({
                "timestamp": timestamp,
                "sender": sender.strip(),
                "content": message.strip(),
                "type": "text"
            })
    else:
        # 通用格式：按行分割
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                # 尝试提取发送者和内容
                if ':' in line or '：' in line:
                    parts = re.split(r'[:：]', line, 1)
                    if len(parts) == 2:
                        segments.append({
                            "sender": parts[0].strip(),
                            "content": parts[1].strip(),
                            "type": "text"
                        })
                    else:
                        segments.append({
                            "content": line,
                            "type": "text"
                        })
                else:
                    segments.append({
                        "content": line,
                        "type": "text"
                    })
    
    return {
        "content": content,
        "segments": segments,
        "format": "chat",
        "total_messages": len(segments)
    }


async def parse_document(file_path: str) -> Dict[str, Any]:
    """
    解析文档文件（PDF、Word、TXT、Markdown）
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == '.pdf':
        return await parse_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return await parse_docx(file_path)
    elif ext in ['.txt', '.md', '.markdown']:
        return await parse_text(file_path)
    else:
        # 默认作为文本读取
        return await parse_text(file_path)


async def parse_pdf(file_path: str) -> Dict[str, Any]:
    """解析 PDF 文件"""
    try:
        import PyPDF2
        
        text_content = []
        segments = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text_content.append(page_text)
                
                segments.append({
                    "page": page_num + 1,
                    "content": page_text,
                    "type": "page"
                })
        
        return {
            "content": "\n\n".join(text_content),
            "segments": segments,
            "format": "pdf",
            "total_pages": len(pdf_reader.pages)
        }
    
    except ImportError:
        # 如果 PyPDF2 未安装，尝试作为文本读取
        return await parse_text(file_path)


async def parse_docx(file_path: str) -> Dict[str, Any]:
    """解析 Word 文档"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        paragraphs = []
        segments = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                segments.append({
                    "content": text,
                    "type": "paragraph"
                })
        
        return {
            "content": "\n\n".join(paragraphs),
            "segments": segments,
            "format": "docx",
            "total_paragraphs": len(paragraphs)
        }
    
    except ImportError:
        return await parse_text(file_path)


async def parse_text(file_path: str) -> Dict[str, Any]:
    """解析纯文本文件"""
    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = await f.read()
    
    # 按段落分割
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    
    segments = [
        {"content": para, "type": "paragraph"}
        for para in paragraphs
    ]
    
    return {
        "content": content,
        "segments": segments,
        "format": "text",
        "total_paragraphs": len(paragraphs)
    }


async def parse_social_media(file_path: str) -> Dict[str, Any]:
    """
    解析社交媒体数据（JSON 格式）
    """
    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = await f.read()
    
    try:
        # 尝试作为 JSON 解析
        data = json.loads(content)
        
        segments = []
        
        if isinstance(data, list):
            # 推文列表
            for item in data:
                if isinstance(item, dict):
                    segments.append({
                        "content": item.get("text", item.get("content", "")),
                        "timestamp": item.get("created_at", item.get("timestamp")),
                        "likes": item.get("likes", item.get("favorite_count", 0)),
                        "type": "post"
                    })
        elif isinstance(data, dict):
            # 单条记录或包含 posts 的包装
            posts = data.get("posts", data.get("tweets", []))
            if posts:
                for item in posts:
                    segments.append({
                        "content": item.get("text", item.get("content", "")),
                        "timestamp": item.get("created_at", item.get("timestamp")),
                        "likes": item.get("likes", item.get("favorite_count", 0)),
                        "type": "post"
                    })
            else:
                segments.append({
                    "content": data.get("text", ""),
                    "type": "post"
                })
        
        return {
            "content": content,
            "segments": segments,
            "format": "social_media",
            "total_posts": len(segments)
        }
    
    except json.JSONDecodeError:
        # 不是 JSON，作为文本处理
        return await parse_text(file_path)


async def parse_audio_transcript(file_path: str) -> Dict[str, Any]:
    """
    解析音视频转录文本
    """
    # 目前作为文本处理，未来可以集成 Whisper API
    return await parse_text(file_path)


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    将长文本分割成块
    """
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # 尽量在句子边界分割
        if end < len(text):
            # 找最近的句号、问号、感叹号
            for punct in ['。', '？', '！', '. ', '? ', '! ']:
                last_punct = chunk.rfind(punct)
                if last_punct > chunk_size * 0.5:  # 至少保留一半内容
                    chunk = chunk[:last_punct + 1]
                    break
        
        chunks.append(chunk.strip())
        start = end - overlap
    
    return chunks


def parse_file(file_path: str, mime_type: str = "") -> str:
    """
    通用文件解析函数 - 同步版本
    根据文件类型调用相应的解析器，返回纯文本内容
    """
    from pathlib import Path
    
    path = Path(file_path)
    ext = path.suffix.lower()
    
    # 简单同步读取，主要用于文本文件
    try:
        if ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    return "\n\n".join(text_content)
            except ImportError:
                pass
        
        elif ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                return "\n\n".join(paragraphs)
            except ImportError:
                pass
        
        # 默认作为文本读取
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"[Parser] Error parsing file {file_path}: {e}")
        return ""
